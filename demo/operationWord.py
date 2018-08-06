#!/usr/bin/env python
# encoding: utf-8
# @author: 曹会金
# @license: (C) Copyright 2013-2017, Node Supply Chain Manager Corporation Limited.
# @contact: deamoncao100@gmail.com
# @software: garner
# @file: operationWord.py.py
# @time: 2018/8/6 11:00
# @desc:  使用python 操作word文档

import docx


# 文件路径
resourcesPath = "..\\resources\\demoFile\\doc\\"
# 对象实例
document = docx.Document()

paragraph = document.add_paragraph("作者: 苏轼", "Subtitle")

prior_paragraph = paragraph.insert_paragraph_before("水调歌头", "Title")

content = """
丙辰中秋，欢饮达旦，大醉，作此篇，兼怀子由。
明月几时有？把酒问青天。
不知天上宫阙，今夕是何年。
我欲乘风归去，又恐琼楼玉宇，高处不胜寒。
起舞弄清影，何似在人间。
转朱阁，低绮户，照无眠。
不应有恨，何事长向别时圆？
人有悲欢离合，月有阴晴圆缺，此事古难全。
但愿人长久，千里共婵娟。
"""

document.add_paragraph(content, "Body Text 2")

document.add_heading("heading")

document.save(resourcesPath + "text.docx")




